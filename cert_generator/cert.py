#!/usr/bin/env python3
"""
Certificate Generator — Rani Durgavati Vishwavidyalaya
=======================================================
Each course has its own template (.doc / .docx).
Only student-specific particulars change per certificate:
  Roll No · Name (English + KrutiDev Hindi) · College · Year · Division · Gender pronoun

How fonts are handled
---------------------
Replacements are made by writing directly into <w:t> XML text nodes.
The <w:rPr> element (which carries font, size, bold, colour) on every
run is NEVER touched — so whatever font the template author set on a run
stays exactly as-is after replacement.

KrutiDev note
-------------
The template runs already carry the correct KrutiDev font (e.g.
"Kruti Dev Display 490", "Kruti Dev 050 Wide").  The student's KrutiDev
name from the Excel column student_name_hindi_KrutiDev is placed into
the matching run and rendered by that font automatically.
Hindi-font runs are NEVER replaced with English text.

.doc support
------------
Loads via LibreOffice (soffice) auto-conversion on the fly.

Excel columns used  (Book1_krutidev.xlsx)
-----------------------------------------
  roll_number · student_name · student_name_hindi_KrutiDev
  college_dec · year_term_code · division · gender
  Lookup by:  roll_number  |  applicationno  |  enroll_number

Setup:  pip install python-docx openpyxl
"""

import sys, os, json, shutil, tempfile, subprocess
from pathlib import Path

try:
    import tkinter as tk
    from tkinter import ttk, messagebox
    from tkinter.filedialog import askopenfilename, asksaveasfilename
except ImportError:
    print("ERROR: tkinter not available. Reinstall Python and tick 'tcl/tk and IDLE'.")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    messagebox.showerror("Missing", "python-docx not installed.\nRun: pip install python-docx")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    messagebox.showerror("Missing", "openpyxl not installed.\nRun: pip install openpyxl")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────────────────────
# Date helpers
# ─────────────────────────────────────────────────────────────────────────────
MONTHS_EN = ["January","February","March","April","May","June",
             "July","August","September","October","November","December"]

# KrutiDev-encoded month names (Kruti Dev 010 font)
MONTHS_KD = ["tuojh","Qjojh","ekpZ","vizSy","ebZ","twu",
             "tqykbZ","vxLr","flrEcj","vDVwcj","uoEcj","fnlEcj"]

def _ordinal(n: int) -> str:
    if 11 <= n <= 13: return f"{n}th"
    return f"{n}{['th','st','nd','rd','th'][min(n % 10, 4)]}"

def build_date_en(day: int, month: int, year: int) -> str:
    """e.g. 02nd November, 2012."""
    return f"{_ordinal(day)} {MONTHS_EN[month-1]}, {year}."

def build_date_kd(day: int, month: int, year: int) -> str:
    """KrutiDev date  e.g.  02 uoEcj] 2012-"""
    return f"{day:02d} {MONTHS_KD[month-1]}] {year}-"


# ─────────────────────────────────────────────────────────────────────────────
# Division lookup  (matches exact values from your Excel)
# ─────────────────────────────────────────────────────────────────────────────
_DIV = {
    "FIRST":       ("First",       "izFke"),
    "SECOND":      ("Second",      "f}rh;"),
    "THIRD":       ("Third",       "r`rh;"),
    "PASS":        ("Pass",        "ikl"),
    "DISTINCT":    ("Distinction", 'fo\'ks"k'),
    "DISTINCTION": ("Distinction", 'fo\'ks"k'),
}

def _norm_div(raw):
    c = raw.strip().upper()
    if c in _DIV: return c
    if c.startswith("FIRST")  or c in ("1ST","I","1"):   return "FIRST"
    if c.startswith("SECOND") or c in ("2ND","II","2"):  return "SECOND"
    if c.startswith("THIRD")  or c in ("3RD","III","3"): return "THIRD"
    if c.startswith("DIST"):                              return "DISTINCT"
    if c.startswith("PASS")   or c == "P":               return "PASS"
    return "FIRST"


# ─────────────────────────────────────────────────────────────────────────────
# Excel loader
# ─────────────────────────────────────────────────────────────────────────────
def load_excel(path):
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    rows = list(wb.active.iter_rows(values_only=True))
    wb.close()
    if not rows:
        raise ValueError("Excel file is empty.")

    hdrs = [str(h).strip() if h is not None else "" for h in rows[0]]
    hup  = [h.upper() for h in hdrs]

    def col(name):
        try:    return hdrs[hup.index(name.upper())]
        except: return None

    rc = col("roll_number")
    if not rc:
        raise ValueError("No 'roll_number' column.\nFound: " + str([h for h in hdrs if h][:15]))

    by_roll={};  by_app={}; by_enr={}
    for row in rows[1:]:
        if not any(row): continue
        rec = {hdrs[i]: (str(v).strip() if v is not None else "")
               for i, v in enumerate(row) if i < len(hdrs)}
        r = rec.get(rc, "").strip().upper()
        if r: by_roll[r] = rec
        a = rec.get(col("applicationno") or "", "").strip().upper()
        if a: by_app[a] = rec
        e = rec.get(col("enroll_number") or "", "").strip().upper()
        if e: by_enr[e] = rec

    return by_roll, by_app, by_enr


def extract_fields(rec):
    """Pull the 7 certificate particulars out of one Excel row."""
    ru = {k.upper(): v for k, v in rec.items()}
    def g(*keys):
        for k in keys:
            v = ru.get(k.upper(), "")
            if v: return v
        return ""

    div_key  = _norm_div(g("division", "div"))
    div_en, div_kd = _DIV.get(div_key, ("First", "izFke"))
    gender   = g("gender", "sex").upper()
    female   = gender in ("F", "FEMALE")

    # ── English gender words ──────────────────────────────────────────────────
    gp = "her" if female else "him"

    # ── KrutiDev gender words ─────────────────────────────────────────────────
    # Pronoun:  bUgsa = उन्हें (him/her)  →  both pronouns use the same word,
    #           but the template uses bUgsa for male; same word works for female too.
    gp_kd    = "mUgsa" if female else "bUgsa"

    # के / की  — genitive particle after college name
    #   male:   "ds" = के  (ke)
    #   female: "dh" = की  (ki)
    ke_ki_kd = "dh" if female else "ds"

    # छात्र / छात्रा  — student (appears at end of college line)
    #   male:   "Nk=" = छात्र
    #   female: "Nk=k" = छात्रा
    student_kd = "Nk=k" if female else "Nk="

    # Sentence ending on the "passed in ... Division" line
    #   male:   "mÙkh.kZ dhA" = उत्तीर्ण किया  (kiya — masculine past tense)
    #   female: "mÙkh.kZ dh"  = उत्तीर्ण की    (ki   — feminine past tense)
    pass_suffix_kd = "mÙkh.kZ dh" if female else "mÙkh.kZ dhA"

    # "certified is done that" line
    #   male:   "tkrk" = जाता  (jaataa — masculine)
    #   female: "tkrh" = जाती  (jaati  — feminine)
    certify_kd = "tkrh" if female else "tkrk"

    return {
        "ROLL_NO":          g("roll_number"),
        "NAME_EN":          g("student_name"),
        "NAME_KD":          g("student_name_hindi_KrutiDev", "student_name_hindi"),
        "COLLEGE":          g("college_dec", "college_name"),
        "COLLEGE_KD":       g("college_hindi_KrutiDev", "college_kd"),
        "YEAR":             g("year_term_code", "year"),
        "DIVISION_EN":      div_en,
        "DIVISION_KD":      div_kd,
        "GENDER_PRO":       gp,
        "GENDER_PRO_KD":    gp_kd,
        "KE_KI_KD":         ke_ki_kd,       # के/की after college name
        "STUDENT_KD":       student_kd,     # छात्र/छात्रा
        "PASS_SUFFIX_KD":   pass_suffix_kd, # किया/की at end of division line
        "CERTIFY_KD":       certify_kd,     # जाता/जाती in certification line
        # DATE_EN / DATE_KD are populated by the date picker in the UI
        "DATE_EN":          "",
        "DATE_KD":          "",
    }


# ─────────────────────────────────────────────────────────────────────────────
# .doc → .docx conversion via LibreOffice
# ─────────────────────────────────────────────────────────────────────────────
def convert_doc(path):
    tmp = tempfile.mkdtemp(prefix="cert_conv_")
    try:
        r = subprocess.run(
            ["soffice","--headless","--convert-to","docx","--outdir",tmp,path],
            capture_output=True, text=True, timeout=60)
        if r.returncode != 0:
            raise RuntimeError("LibreOffice error:\n" + (r.stderr or r.stdout))
        out = Path(tmp) / (Path(path).stem + ".docx")
        if not out.exists():
            hits = list(Path(tmp).glob("*.docx"))
            if not hits: raise RuntimeError("No .docx produced after conversion.")
            out = hits[0]
        return str(out)
    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice not found.\n"
            "Install from libreoffice.org\n"
            "or open the .doc in Word and Save As .docx manually.")
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice conversion timed out.")


_W     = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _replace_in_para(para_elem, search, replacement):
    """
    Replace `search` with `replacement` inside one paragraph element,
    handling the case where the search text is split across multiple <w:t> nodes.
    Returns number of replacements made.
    """
    t_nodes = list(para_elem.iter(f"{{{_W}}}t"))
    if not t_nodes:
        return 0

    # Build full concatenated text and track each node's position
    parts     = [(t, t.text or "") for t in t_nodes]
    full_text = "".join(txt for _, txt in parts)

    if search not in full_text:
        return 0

    count  = 0
    offset = 0
    while True:
        idx = full_text.find(search, offset)
        if idx == -1:
            break
        end_idx = idx + len(search)

        # Find which t_nodes overlap the match span [idx, end_idx)
        pos      = 0
        affected = []
        for ni, (t, txt) in enumerate(parts):
            node_end = pos + len(txt)
            if node_end > idx and pos < end_idx:
                rel_s = max(0, idx - pos)
                rel_e = min(len(txt), end_idx - pos)
                affected.append((ni, rel_s, rel_e))
            pos = node_end

        if not affected:
            offset = idx + 1
            continue

        # Write replacement into the first affected node
        ni0, rs0, re0 = affected[0]
        t0, txt0 = parts[ni0]
        new0 = txt0[:rs0] + replacement + txt0[re0:]
        t0.text = new0
        t0.set(_SPACE, "preserve")
        parts[ni0] = (t0, new0)

        # Remove the matched portion from every subsequent affected node
        for ni, rs, re in affected[1:]:
            t, txt = parts[ni]
            new = txt[:rs] + txt[re:]
            t.text = new or None
            if new:
                t.set(_SPACE, "preserve")
            parts[ni] = (t, new)

        # Rebuild full_text for the next iteration
        full_text = full_text[:idx] + replacement + full_text[end_idx:]
        offset = idx + len(replacement)
        count += 1

    return count


def _replace_in_element(root_elem, search, replacement):
    """Apply _replace_in_para to every paragraph under root_elem."""
    total = 0
    for p in root_elem.iter(f"{{{_W}}}p"):
        total += _replace_in_para(p, search, replacement)
    return total


def apply_mappings(doc, mappings, values):
    """
    Apply all active mappings to the document.
    `values` is read directly from the UI entry widgets at call time,
    so any manual edits the user made are included.
    Returns total replacement count.
    """
    total = 0
    for m in mappings:
        if not m.get("active", True):
            continue
        search = m["search_text"]
        vk     = m["value_key"]
        if "{" in vk:
            try:    repl = vk.format(**values)
            except: repl = ""
        else:
            repl = values.get(vk, "")

        if not search or not repl or search == repl:
            continue

        total += _replace_in_element(doc.element.body, search, repl)
        for sec in doc.sections:
            for hf in (sec.header, sec.footer,
                       sec.first_page_header, sec.first_page_footer):
                if hf is not None:
                    try: total += _replace_in_element(hf._element, search, repl)
                    except: pass

    return total


# ─────────────────────────────────────────────────────────────────────────────
# Default mappings for the RDV University BA template
# NOTE: The Hindi college line is intentionally omitted — that run uses a
# KrutiDev font; replacing it with an English college name would look garbled.
# If your template has an English college placeholder instead, add a mapping.
# ─────────────────────────────────────────────────────────────────────────────
DEFAULT_MAPPINGS = [
    # Roll number — appears in 4 places (Hindi + English sides of certificate)
    {"label": "Roll No",
     "search_text": "ROLL",
     "value_key":   "ROLL_NO",
     "active": True},

    # Student name English — NuptialScript font run (preserved automatically)
    {"label": "Name (English)",
     "search_text": "NAME_E",
     "value_key":   "NAME_EN",
     "active": True},

    # Student name KrutiDev — Kruti Dev Display 490 font run (preserved)
    # The Excel column student_name_hindi_KrutiDev contains already-encoded KrutiDev text
    {"label": "Name (KrutiDev Hindi)",
     "search_text": "Nk=uke",
     "value_key":   "NAME_KD",
     "active": True},

    # College — English NuptialScript run
    # The Hindi college line is NOT replaced (KrutiDev font, English would look garbled)
    {"label": "College (English)",
     "search_text": "COLLEGE_E",
     "value_key":   "COLLEGE",
     "active": True},

    # College — KrutiDev Hindi (Kruti Dev 050 Wide font)
    # Reads from Excel column: college_hindi_KrutiDev
    # If that column is absent, COLLEGE_KD is blank and this mapping is skipped
     {"label": "College (KrutiDev Hindi)",
     "search_text": "d‚yst",
     "value_key":   "COLLEGE_KD",
     "active": True},

    # Year — Tahoma font run (preserved)
    {"label": "Year",
     "search_text": "YEAR",
     "value_key":   "YEAR",
     "active": True},

    # Division English — trailing space is significant (it's its own run)
    {"label": "Division (English)",
     "search_text": "DIV ",
     "value_key":   "{DIVISION_EN} ",
     "active": True},

    # Division KrutiDev — Kruti Dev Display 490 run, trailing space significant
    {"label": "Division (KrutiDev Hindi)",
     "search_text": "foHkkx",
     "value_key":   "{DIVISION_KD} ",
     "active": True},

    # Gender pronoun English — inside " is being awarded to him."
    {"label": "Gender pronoun (English)",
     "search_text": " him.",
     "value_key":   " {GENDER_PRO}.",
     "active": True},

    # Gender pronoun KrutiDev — "bUgsa" = him, "mUgsa" = her
    {"label": "Gender pronoun (KrutiDev Hindi)",
     "search_text": "bUgsa ",
     "value_key":   "{GENDER_PRO_KD} ",
     "active": True},

    # ke/ki — genitive particle after college name
    # "ds Nk=" = के छात्र (male)  →  "dh Nk=k" = की छात्रा (female)
    # We replace "ds" alone to keep it minimal, and "Nk=" alone for छात्र/छात्रा
    {"label": "ke/ki  (ds / dh)",
     "search_text": "ds Nk=",
     "value_key":   "{KE_KI_KD} {STUDENT_KD}",
     "active": True},

    # Certification verb — "tkrk" = जाता (male)  /  "tkrh" = जाती (female)
    {"label": "Certify verb  (tkrk / tkrh)",
     "search_text": "tkrk gS fd",
     "value_key":   "{CERTIFY_KD} gS fd",
     "active": True},

    # Passed suffix — male: mÙkh.kZ dhA  /  female: mÙkh.kZ dh
    {"label": "Passed suffix  (dhA / dh)",
     "search_text": "mÙkh.kZ dhA",
     "value_key":   "PASS_SUFFIX_KD",
     "active": True},

    # Date — English  e.g.  02nd November, 2012.
    {"label": "Date (English)",
     "search_text": "DATE_E",
     "value_key":   "DATE_EN",
     "active": True},

    # Date — KrutiDev Hindi  e.g.   02 uoEcj] 2012-
    # search_text matches only the date portion (after the " %  " separator)
    {"label": "Date (KrutiDev Hindi)",
     "search_text": "fnukadfnukad",
     "value_key":   "DATE_KD",
     "active": True},
]


# ─────────────────────────────────────────────────────────────────────────────
# UI colours
# ─────────────────────────────────────────────────────────────────────────────
BG    = "#1a1f2e"
PANEL = "#252d3e"
INPUT = "#0f1520"
GOLD  = "#e8c96d"
BLUE  = "#7eb8f7"
GREEN = "#5fba7d"
MUTED = "#4a5568"
RED   = "#e05c5c"
DARK  = "#0d1117"


# ─────────────────────────────────────────────────────────────────────────────
# Application
# ─────────────────────────────────────────────────────────────────────────────
class App(tk.Tk):

    # Fixed field order shown in the UI
    # DATE_EN and DATE_KD are NOT in this list — they have a dedicated date picker
    FIELDS = [
        ("ROLL_NO",        "Roll No"),
        ("NAME_EN",        "Name (English)"),
        ("NAME_KD",        "Name (KrutiDev Hindi)"),
        ("COLLEGE",        "College"),
        ("COLLEGE_KD",     "College (KrutiDev Hindi)"),
        ("YEAR",           "Year of passing"),
        ("DIVISION_EN",    "Division (English)"),
        ("DIVISION_KD",    "Division (KrutiDev Hindi)"),
        ("GENDER_PRO",     "Gender pronoun  him/her"),
        ("GENDER_PRO_KD",  "Gender pronoun  bUgsa/mUgsa"),
        ("KE_KI_KD",       "ke / ki  —  ds / dh"),
        ("STUDENT_KD",     "Student  —  Nk= / Nk=k"),
        ("PASS_SUFFIX_KD", "Passed suffix  —  dhA / dh"),
        ("CERTIFY_KD",     "Certify verb  —  tkrk / tkrh"),
    ]

    def __init__(self):
        super().__init__()
        self.title("Certificate Generator — RDV University")
        self.configure(bg=BG)
        self.geometry("1150x840")
        self.minsize(940, 660)

        self._template_path = ""
        self._db_roll = {};  self._db_app = {};  self._db_enr = {}
        # _data_vars: key -> tk.StringVar, created once in _build_left, never recreated
        self._data_vars: dict[str, tk.StringVar] = {}
        self._mappings = [dict(m) for m in DEFAULT_MAPPINGS]
        self._sel_idx: int | None = None
        # Date section state
        # _date_choice: holds the currently selected dropdown value, either a
        #   preset string like "2nd November, 2012." or the sentinel "Custom…"
        self._date_choice  = tk.StringVar(value="")
        self._date_custom  = tk.StringVar(value="")   # typed when Custom… is chosen
        self._date_kd_var  = tk.StringVar(value="")   # KrutiDev date (always auto)

        self._build_ui()
        self._try_autoload()

    # ── build UI ──────────────────────────────────────────────────────────────

    def _build_ui(self):
        self._build_topbar()
        self._build_excel_bar()
        tk.Frame(self, bg="#2a3040", height=1).pack(fill="x")
        body = tk.Frame(self, bg=BG)
        body.pack(fill="both", expand=True, padx=10, pady=8)
        self._build_left(body)
        self._build_right(body)
        self._sv = tk.StringVar(value="Load a template (.doc/.docx) and Excel file to begin.")
        tk.Label(self, textvariable=self._sv,
                 font=("Courier", 8), bg=DARK, fg="#5588aa",
                 anchor="w", padx=12, pady=5).pack(fill="x", side="bottom")

    def _build_topbar(self):
        bar = tk.Frame(self, bg=DARK, pady=9)
        bar.pack(fill="x")
        tk.Label(bar, text="🎓  Certificate Generator — RDV University",
                 font=("Georgia", 14, "bold"), bg=DARK, fg=GOLD
                 ).pack(side="left", padx=16)
        for txt, cmd in [("📄 Load Template",   self._load_template),
                         ("💾 Save Mappings",    self._save_mappings),
                         ("📋 Load Mappings",    self._load_mappings)]:
            tk.Button(bar, text=txt, font=("Helvetica", 9),
                      bg=PANEL, fg=BLUE, relief="flat", cursor="hand2",
                      padx=9, pady=4, command=cmd).pack(side="left", padx=3)
        self._tmpl_lbl = tk.Label(bar, text="No template loaded",
                                  font=("Courier", 8), bg=DARK, fg=MUTED)
        self._tmpl_lbl.pack(side="left", padx=10)

    def _build_excel_bar(self):
        bar = tk.Frame(self, bg=DARK, pady=4)
        bar.pack(fill="x")
        tk.Label(bar, text="📊 Excel:", font=("Helvetica", 9, "bold"),
                 bg=DARK, fg=BLUE, padx=12).pack(side="left")
        self._xpv = tk.StringVar()
        e = tk.Entry(bar, textvariable=self._xpv, font=("Courier", 9), width=58,
                     bg=INPUT, fg="#d4e8ff", insertbackground="#d4e8ff",
                     relief="flat", bd=5)
        e.pack(side="left", padx=(0,6), ipady=3)
        e.bind("<Return>", lambda _: self._load_excel())
        tk.Button(bar, text="Browse…", font=("Helvetica", 9), bg=PANEL, fg=BLUE,
                  relief="flat", cursor="hand2", padx=8, pady=3,
                  command=self._browse_excel).pack(side="left", padx=2)
        tk.Button(bar, text="Load ↵", font=("Helvetica", 9, "bold"),
                  bg="#2ecc71", fg="#1a1a2e", relief="flat", cursor="hand2",
                  padx=8, pady=3, command=self._load_excel).pack(side="left", padx=2)
        self._xinfo = tk.Label(bar, text="", font=("Courier", 8),
                               bg=DARK, fg=MUTED, padx=8)
        self._xinfo.pack(side="left")

    def _build_left(self, parent):
        # Outer frame that takes up the left side of the body
        outer = tk.Frame(parent, bg=BG)
        outer.pack(side="left", fill="both", expand=True)

        # Scrollable canvas so nothing gets clipped on small screens
        canvas = tk.Canvas(outer, bg=BG, highlightthickness=0)
        vsb    = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        col = tk.Frame(canvas, bg=BG)
        col_id = canvas.create_window((0, 0), window=col, anchor="nw")

        def _on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def _on_canvas_configure(event):
            canvas.itemconfig(col_id, width=event.width)
        col.bind("<Configure>", _on_frame_configure)
        canvas.bind("<Configure>", _on_canvas_configure)

        # Mouse-wheel scrolling (Windows + Linux + macOS)
        def _on_mousewheel(event):
            if event.num == 4:
                canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                canvas.yview_scroll(1, "units")
            else:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        canvas.bind_all("<Button-4>",   _on_mousewheel)
        canvas.bind_all("<Button-5>",   _on_mousewheel)

        # Lookup bar
        lf = tk.Frame(col, bg=DARK, bd=1, relief="groove", pady=8)
        lf.pack(fill="x")
        tk.Label(lf, text="Student Lookup",
                 font=("Helvetica", 10, "bold"), bg=DARK, fg=GOLD, padx=10
                 ).pack(anchor="w")
        tk.Label(lf, text="  Accepts:  Roll Number  ·  Application No  ·  Enrolment No",
                 font=("Courier", 7), bg=DARK, fg=MUTED).pack(anchor="w")
        row = tk.Frame(lf, bg=DARK)
        row.pack(fill="x", padx=10, pady=6)
        tk.Label(row, text="ID:", font=("Helvetica", 9, "bold"),
                 bg=DARK, fg="#9ab").pack(side="left")
        self._idv = tk.StringVar()
        e = tk.Entry(row, textvariable=self._idv, font=("Courier", 14), width=18,
                     bg=INPUT, fg=GOLD, insertbackground=GOLD, relief="flat", bd=6)
        e.pack(side="left", padx=6)
        e.bind("<Return>", lambda _: self._lookup())
        tk.Button(row, text="Look Up", font=("Helvetica", 9, "bold"),
                  bg=GOLD, fg="#1a1a2e", relief="flat", cursor="hand2",
                  command=self._lookup).pack(side="left")
        self._hint = tk.Label(lf, text="  Load Excel to see available IDs",
                              font=("Courier", 7), bg=DARK, fg=MUTED)
        self._hint.pack(anchor="w", padx=10, pady=(0,4))

        # Certificate fields — created ONCE here; lookup only calls var.set()
        df = tk.Frame(col, bg=DARK, bd=1, relief="groove")
        df.pack(fill="x", pady=(8,0))
        tk.Label(df, text="Certificate Particulars  —  edit freely before generating",
                 font=("Helvetica", 9, "bold"), bg=DARK, fg=BLUE,
                 pady=6, padx=10).pack(anchor="w")
        tk.Label(df,
                 text="Only these fields change between students.  "
                      "All template fonts, sizes and colours are preserved automatically.",
                 font=("Helvetica", 7), bg=DARK, fg=MUTED,
                 padx=10, wraplength=540, justify="left").pack(anchor="w", pady=(0,6))

        grid = tk.Frame(df, bg=DARK, padx=12, pady=4)
        grid.pack(fill="x")
        for row_i, (key, label) in enumerate(self.FIELDS):
            tk.Label(grid, text=label + ":", font=("Helvetica", 9),
                     bg=DARK, fg=BLUE, anchor="e", width=28
                     ).grid(row=row_i, column=0, sticky="e", pady=4, padx=(0,8))
            var = tk.StringVar()
            self._data_vars[key] = var          # store reference — NEVER recreate
            tk.Entry(grid, textvariable=var, font=("Courier", 10), width=36,
                     bg=INPUT, fg="#d4e8ff", insertbackground="#d4e8ff",
                     relief="flat", bd=4
                     ).grid(row=row_i, column=1, sticky="ew", pady=4)
        grid.columnconfigure(1, weight=1)

        # Date picker
        self._build_date_section(col)

        # Action buttons
        acts = tk.Frame(col, bg=BG, pady=8)
        acts.pack(fill="x")
        for lbl, bg_, abg, cmd in [
            ("🎓  Generate DOCX", "#2ecc71", "#27ae60", self._generate),
            ("🖨️   Print",          "#e67e22", "#d35400", self._print),
            ("💾  Save As…",       "#9b59b6", "#8e44ad", self._save_as)]:
            tk.Button(acts, text=lbl, font=("Helvetica", 9, "bold"),
                      bg=bg_, fg="white", activebackground=abg,
                      relief="flat", cursor="hand2", padx=14, pady=7,
                      command=cmd).pack(side="left", padx=4)

    # ── Preset date list (English) — edit freely to match your convocation schedule ──
    # Format must match the template exactly, e.g. "2nd November, 2012."
    # The KrutiDev date is always auto-converted from this English date.
    PRESET_DATES = [
        "29th December, 1964.",
        "29th January, 1969.",
        "07th February, 1978.",
        "01st February, 1983.",
        "07th January, 1984.",
        "28th February, 2002.",
        "07th March, 2005.",
        "12th October, 2006.",
        "09th October, 2007.",
        "14th October, 2008.",
        "24th October, 2009.",
        "24th June, 2010.",
        "24th June, 2011.",
        "02nd November, 2012.",
        "30th December, 2013.",
        "30th October, 2014.",
        "06th March, 2017.",
        "03rd May, 2018.",
        "14th March, 2019.",
        "06th February, 2021.",
        "12th January, 2022.",
        "29th December, 2022.",
        "05th October, 2023.",
    ]
    _CUSTOM_SENTINEL = "✏  Type custom date…"

    def _build_date_section(self, parent):
        """
        Date section:
          • A dropdown listing PRESET_DATES plus a "Type custom date…" option at the end
          • When the custom option is chosen a text box slides into view for free typing
          • KrutiDev date is always auto-converted from the English date string

        DATE_EN and DATE_KD are registered into self._data_vars so _get_values()
        picks them up without any special casing.
        """
        for key in ("DATE_EN", "DATE_KD"):
            if key not in self._data_vars:
                self._data_vars[key] = tk.StringVar()

        outer = tk.Frame(parent, bg=DARK, bd=1, relief="groove")
        outer.pack(fill="x", pady=(8,0))

        hdr = tk.Frame(outer, bg=DARK)
        hdr.pack(fill="x", padx=10, pady=(6,2))
        tk.Label(hdr, text="📅  Certificate Date",
                 font=("Helvetica", 9, "bold"), bg=DARK, fg=BLUE).pack(side="left")
        tk.Label(hdr,
                 text='  -- select a preset, or choose "Type custom date..." at the bottom',
                 font=("Helvetica", 7), bg=DARK, fg=MUTED).pack(side="left")

        # ── Preset dropdown ───────────────────────────────────────────
        pick_row = tk.Frame(outer, bg=DARK, padx=12, pady=4)
        pick_row.pack(fill="x")

        tk.Label(pick_row, text="Date:", font=("Helvetica", 9),
                 bg=DARK, fg=BLUE, width=6, anchor="e").pack(side="left", padx=(0,6))

        dropdown_values = self.PRESET_DATES + [self._CUSTOM_SENTINEL]
        # Initialise choice to first preset
        if not self._date_choice.get():
            self._date_choice.set(self.PRESET_DATES[0])

        self._date_cb = ttk.Combobox(
            pick_row,
            textvariable=self._date_choice,
            values=dropdown_values,
            width=34,
            state="readonly",
            font=("Courier", 10))
        self._date_cb.pack(side="left")

        # ── Custom text box (hidden until sentinel chosen) ─────────────
        self._date_custom_row = tk.Frame(outer, bg=DARK, padx=12)
        # not packed yet — shown only when Custom… is selected

        tk.Label(self._date_custom_row, text="Custom date:",
                 font=("Helvetica", 9), bg=DARK, fg=GOLD,
                 width=13, anchor="e").pack(side="left")
        self._date_custom_entry = tk.Entry(
            self._date_custom_row,
            textvariable=self._date_custom,
            font=("Courier", 10), width=34,
            bg=INPUT, fg=GOLD, insertbackground=GOLD,
            relief="flat", bd=4)
        self._date_custom_entry.pack(side="left", padx=(6,0))

        # ── KrutiDev preview ──────────────────────────────────────────
        prev_row = tk.Frame(outer, bg=DARK)
        prev_row.pack(fill="x", padx=12, pady=(2,8))
        tk.Label(prev_row, text="Hindi date:", font=("Helvetica", 8),
                 bg=DARK, fg=MUTED).pack(side="left")
        self._date_kd_label = tk.Label(prev_row, textvariable=self._date_kd_var,
                                       font=("Courier", 9), bg=DARK, fg=GREEN)
        self._date_kd_label.pack(side="left", padx=6)

        # Bind changes
        self._date_choice.trace_add("write", lambda *_: self._on_date_change())
        self._date_custom.trace_add("write", lambda *_: self._on_date_change())

        self._on_date_change()   # populate on startup

    def _on_date_change(self, *_):
        """
        Called whenever the dropdown selection or custom text changes.
        Parses the chosen English date string → builds KrutiDev equivalent.
        Shows/hides the custom entry row as needed.
        """
        import re
        choice = self._date_choice.get()

        if choice == self._CUSTOM_SENTINEL:
            # Show custom entry
            self._date_custom_row.pack(fill="x", padx=12, pady=(0,4),
                                       before=self._date_kd_label.master)
            self._date_custom_entry.focus_set()
            en = self._date_custom.get().strip()
        else:
            # Hide custom entry
            self._date_custom_row.pack_forget()
            en = choice

        # Auto-build KrutiDev date by parsing the English date string
        # Expected format: "2nd November, 2012."  or similar
        kd = ""
        if en:
            m = re.match(
                r"(\d{1,2})(?:st|nd|rd|th)\s+(\w+),?\s*(\d{4})",
                en, re.IGNORECASE)
            if m:
                day_s, mon_s, yr_s = m.group(1), m.group(2), m.group(3)
                mon_s_cap = mon_s.capitalize()
                if mon_s_cap in MONTHS_EN:
                    month = MONTHS_EN.index(mon_s_cap) + 1
                    kd = build_date_kd(int(day_s), month, int(yr_s))

        self._date_kd_var.set(kd if kd else "(type a date like  5th March, 2024.)")

        # Push into data_vars
        if "DATE_EN" in self._data_vars:
            self._data_vars["DATE_EN"].set(en)
        if "DATE_KD" in self._data_vars:
            self._data_vars["DATE_KD"].set(kd)

    def _build_right(self, parent):
        sb = tk.Frame(parent, bg=PANEL, width=380)
        sb.pack(side="right", fill="y", padx=(10,0))
        sb.pack_propagate(False)

        tk.Label(sb, text="🔁  Search → Replace Mappings",
                 font=("Helvetica", 10, "bold"), bg=PANEL, fg=BLUE,
                 pady=8).pack(fill="x", padx=10)
        tk.Label(sb,
                 text="Left = exact text in your template\n"
                      "Right = data key  (or {KEY} pattern)\n"
                      "Fonts are preserved — only the text changes.",
                 font=("Helvetica", 7), bg=PANEL, fg=MUTED,
                 wraplength=350, justify="left").pack(fill="x", padx=10, pady=(0,6))
        tk.Frame(sb, bg="#3a4460", height=1).pack(fill="x", padx=10, pady=(0,4))

        # Listbox
        lf2 = tk.Frame(sb, bg=PANEL)
        lf2.pack(fill="x", padx=10)
        self._lb = tk.Listbox(lf2, font=("Courier", 8), height=10,
                              bg=INPUT, fg="#aaccee",
                              selectbackground="#2a4a6a", selectforeground="white",
                              relief="flat", bd=0, activestyle="none")
        self._lb.pack(fill="x")
        self._lb.bind("<<ListboxSelect>>", self._on_sel)

        # Listbox buttons
        br = tk.Frame(sb, bg=PANEL)
        br.pack(fill="x", padx=10, pady=(4,0))
        tk.Button(br, text="+ Add",    font=("Helvetica",8,"bold"), bg=GREEN,
                  fg="#1a1a2e", relief="flat", cursor="hand2",
                  command=self._add_map).pack(side="left", padx=(0,4))
        tk.Button(br, text="✕ Remove", font=("Helvetica",8), bg=RED,
                  fg="white", relief="flat", cursor="hand2",
                  command=self._rm_map).pack(side="left")
        tk.Button(br, text="↑", font=("Helvetica",8,"bold"), width=2,
                  bg=PANEL, fg=BLUE, relief="flat", cursor="hand2",
                  command=lambda: self._mv_map(-1)).pack(side="right", padx=1)
        tk.Button(br, text="↓", font=("Helvetica",8,"bold"), width=2,
                  bg=PANEL, fg=BLUE, relief="flat", cursor="hand2",
                  command=lambda: self._mv_map(1)).pack(side="right")

        # Active toggle
        self._active_var = tk.BooleanVar(value=True)
        tk.Checkbutton(sb, text="Active (uncheck to skip this mapping)",
                       variable=self._active_var, font=("Helvetica",8),
                       bg=PANEL, fg=MUTED, selectcolor=INPUT, activebackground=PANEL,
                       command=self._toggle_active).pack(anchor="w", padx=12, pady=(6,0))

        # Edit area
        tk.Frame(sb, bg="#3a4460", height=1).pack(fill="x", padx=10, pady=(8,4))
        tk.Label(sb, text="Edit Selected Mapping",
                 font=("Helvetica",9,"bold"), bg=PANEL, fg=BLUE
                 ).pack(fill="x", padx=10)
        ep = tk.Frame(sb, bg=PANEL, padx=10)
        ep.pack(fill="x", pady=4)
        self._ev = {}
        for field in ("label", "search_text", "value_key"):
            row = tk.Frame(ep, bg=PANEL, pady=2)
            row.pack(fill="x")
            tk.Label(row, text=field.replace("_"," ").title()+":",
                     width=12, anchor="w", font=("Helvetica",8),
                     bg=PANEL, fg=MUTED).pack(side="left")
            var = tk.StringVar(); self._ev[field] = var
            tk.Entry(row, textvariable=var, font=("Courier",9),
                     bg=INPUT, fg="#cde", insertbackground="#cde",
                     relief="flat", bd=4).pack(side="left", fill="x", expand=True)

        tk.Button(sb, text="✓ Apply Changes", font=("Helvetica",9,"bold"),
                  bg=GOLD, fg="#1a1a2e", relief="flat", cursor="hand2", pady=6,
                  command=self._apply_edit).pack(fill="x", padx=10, pady=(6,0))
        tk.Button(sb, text="↺ Reset to Defaults", font=("Helvetica",8),
                  bg=PANEL, fg=MUTED, relief="flat", cursor="hand2",
                  command=self._reset_maps).pack(fill="x", padx=10, pady=(4,0))

        self._refresh_lb()

    # ── autoload ──────────────────────────────────────────────────────────────

    def _try_autoload(self):
        here = Path(__file__).parent
        for name in ["Book1_krutidev.xlsx","students.xlsx","data.xlsx"]:
            p = here/name
            if p.exists():
                self._xpv.set(str(p))
                self._status(f"Found {name} — press Load ↵ or Enter.")
                break

    # ── template ──────────────────────────────────────────────────────────────

    def _load_template(self):
        path = askopenfilename(title="Load Certificate Template",
                               filetypes=[("Word Documents","*.docx *.doc"),
                                          ("All files","*.*")])
        if not path: return
        actual = path
        if path.lower().endswith(".doc"):
            self._status("⏳  Converting .doc → .docx via LibreOffice…")
            self.update_idletasks()
            try:    actual = convert_doc(path)
            except RuntimeError as e:
                messagebox.showerror("Conversion Failed", str(e)); return
        try:
            Document(actual).paragraphs   # quick parse test
        except Exception as e:
            messagebox.showerror("Template Error", f"Cannot open:\n{e}"); return
        self._template_path = actual
        sfx = "  (converted from .doc)" if path.lower().endswith(".doc") else ""
        self._tmpl_lbl.configure(text=f"✓  {Path(path).name}{sfx}", fg=GREEN)
        self._status(f"Template ready: {Path(path).name}{sfx}")

    # ── excel ─────────────────────────────────────────────────────────────────

    def _browse_excel(self):
        p = askopenfilename(title="Select Excel",
                            filetypes=[("Excel","*.xlsx *.xlsm *.xls"),("All","*.*")])
        if p: self._xpv.set(p); self._load_excel()

    def _load_excel(self):
        raw = self._xpv.get().strip().strip('"')
        if not raw: messagebox.showwarning("No path","Enter the Excel file path."); return
        path = os.path.normpath(raw)
        if not os.path.isfile(path):
            messagebox.showerror("Not found", f"Cannot find:\n  {path}"); return
        self._status("⏳  Loading Excel…"); self.update_idletasks()
        try:    r, a, e = load_excel(path)
        except Exception as ex:
            messagebox.showerror("Excel Error", str(ex)); return
        if not r:
            messagebox.showwarning("No data","No student records found."); return
        self._db_roll=r; self._db_app=a; self._db_enr=e
        sample = "  ".join(list(r.keys())[:5])
        self._hint.configure(
            text=f"  Sample rolls: {sample}{'  …' if len(r)>5 else ''}",
            fg=GREEN)
        self._xinfo.configure(
            text=f"✓  {Path(path).name}  ({len(r):,} students)", fg=GREEN)
        self._status(f"✅  Loaded {len(r):,} students from {Path(path).name}")

    # ── lookup ────────────────────────────────────────────────────────────────

    def _lookup(self):
        if not self._db_roll:
            messagebox.showwarning("No Excel","Load an Excel file first."); return
        q = self._idv.get().strip().upper()
        if not q: return
        rec = (self._db_roll.get(q) or self._db_app.get(q) or self._db_enr.get(q))
        if not rec:
            sample = ", ".join(list(self._db_roll.keys())[:8])
            messagebox.showwarning("Not found",
                f"No record for '{q}'.\n\nSample roll numbers:\n{sample}"); return

        fields = extract_fields(rec)
        # Update each StringVar in place — the Entry widgets stay bound, edits work
        for key, var in self._data_vars.items():
            var.set(fields.get(key, ""))

        self._status(
            f"✅  {fields['NAME_EN']}  |  {rec.get('exampassed','')}  |  "
            f"{fields['DIVISION_EN']} Division  |  {fields['COLLEGE'][:40]}")

    # ── mappings panel ────────────────────────────────────────────────────────

    def _refresh_lb(self):
        self._lb.delete(0, tk.END)
        for m in self._mappings:
            tick = "✓" if m.get("active",True) else "✗"
            self._lb.insert(tk.END,
                f"  {tick} {m['search_text'][:24]:24s}  →  {m['value_key'][:20]}")

    def _on_sel(self, _):
        sel = self._lb.curselection()
        if not sel: return
        self._sel_idx = sel[0]
        m = self._mappings[self._sel_idx]
        self._ev["label"].set(m.get("label",""))
        self._ev["search_text"].set(m.get("search_text",""))
        self._ev["value_key"].set(m.get("value_key",""))
        self._active_var.set(m.get("active", True))

    def _toggle_active(self):
        if self._sel_idx is None: return
        self._mappings[self._sel_idx]["active"] = self._active_var.get()
        self._refresh_lb()

    def _apply_edit(self):
        if self._sel_idx is None:
            messagebox.showwarning("Nothing selected","Click a mapping row first."); return
        self._mappings[self._sel_idx].update({
            "label":       self._ev["label"].get(),
            "search_text": self._ev["search_text"].get(),
            "value_key":   self._ev["value_key"].get(),
        })
        self._refresh_lb()
        self._status("Mapping updated.")

    def _add_map(self):
        m = {"label":"New","search_text":"find this","value_key":"VALUE_KEY","active":True}
        self._mappings.append(m); self._refresh_lb()
        idx = len(self._mappings)-1
        self._lb.selection_clear(0,tk.END); self._lb.selection_set(idx)
        self._sel_idx = idx; self._on_sel(None)

    def _rm_map(self):
        if self._sel_idx is None: return
        self._mappings.pop(self._sel_idx); self._sel_idx=None; self._refresh_lb()

    def _mv_map(self, d):
        if self._sel_idx is None: return
        ni = self._sel_idx + d
        if ni < 0 or ni >= len(self._mappings): return
        self._mappings[self._sel_idx], self._mappings[ni] = \
            self._mappings[ni], self._mappings[self._sel_idx]
        self._sel_idx = ni; self._refresh_lb()
        self._lb.selection_clear(0,tk.END); self._lb.selection_set(ni)

    def _reset_maps(self):
        self._mappings = [dict(m) for m in DEFAULT_MAPPINGS]
        self._refresh_lb(); self._status("Mappings reset to defaults.")

    def _save_mappings(self):
        path = asksaveasfilename(defaultextension=".certmaps",
            filetypes=[("Cert Mappings","*.certmaps"),("JSON","*.json")])
        if path:
            with open(path,"w",encoding="utf-8") as f:
                json.dump(self._mappings, f, indent=2, ensure_ascii=False)
            self._status(f"💾  Saved: {path}")

    def _load_mappings(self):
        path = askopenfilename(
            filetypes=[("Cert Mappings","*.certmaps *.json"),("All","*.*")])
        if path:
            with open(path,encoding="utf-8") as f:
                self._mappings = json.load(f)
            self._refresh_lb(); self._status(f"Loaded {len(self._mappings)} rules")

    # ── generate ──────────────────────────────────────────────────────────────

    def _get_values(self):
        """
        Read current values directly from the tk.StringVar objects tied to
        the Entry widgets.  Any manual edits the user made ARE included here.
        """
        return {key: var.get() for key, var in self._data_vars.items()}

    def _generate(self):
        if not self._template_path or not os.path.isfile(self._template_path):
            messagebox.showwarning("No Template",
                "Click 📄 Load Template to choose a .doc or .docx file."); return None
        values = self._get_values()
        if not any(values.get(k,"") for k in ("ROLL_NO","NAME_EN","NAME_KD")):
            messagebox.showwarning("No Data",
                "Look up a student first, or type values into the fields."); return None

        rid = values.get("ROLL_NO","cert")
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx",
                                           prefix=f"cert_{rid}_")
        tmp.close()
        self._status("⏳  Generating…"); self.update_idletasks()
        try:
            doc   = Document(self._template_path)
            count = apply_mappings(doc, self._mappings, values)
            doc.save(tmp.name)
        except Exception as e:
            messagebox.showerror("Error", str(e)); return None
        self._open_file(tmp.name)
        self._status(f"✅  {count} replacements made  →  {tmp.name}")
        return tmp.name

    def _open_file(self, path):
        try:
            if sys.platform == "win32":    os.startfile(path)                    # type: ignore
            elif sys.platform == "darwin": subprocess.Popen(["open", path])
            else:                          subprocess.Popen(["xdg-open", path])
        except Exception as e:
            self._status(f"Saved: {path}  (auto-open failed: {e})")

    def _print(self):
        path = self._generate()
        if not path: return
        try:
            if sys.platform == "win32": os.startfile(path, "print")              # type: ignore
            else:                       subprocess.Popen(["lpr", path])
            self._status(f"🖨️  Sent to printer: {path}")
        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    def _save_as(self):
        path = self._generate()
        if not path: return
        rid  = self._data_vars["ROLL_NO"].get() or "output"
        dest = asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document","*.docx"),("All","*.*")],
            initialfile=f"certificate_{rid}.docx")
        if dest:
            shutil.copy2(path, dest)
            self._status(f"💾  Saved: {dest}")

    def _status(self, msg):
        self._sv.set(msg); self.update_idletasks()


if __name__ == "__main__":
    App().mainloop()